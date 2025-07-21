"""
Service de génération de PDF pour les baux avec architecture SOLID
"""
from abc import ABC, abstractmethod
from typing import Dict, Any, Optional, List
import os
import io
from datetime import datetime, date
from pathlib import Path
import hashlib
import json

# Imports pour la génération PDF
try:
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.lib.colors import black, blue, red
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.platypus.tableofcontents import TableOfContents
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# Imports pour la conversion Word (optionnel)
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


class ILeaseDocumentGenerator(ABC):
    """Interface pour les générateurs de documents de bail"""
    
    @abstractmethod
    def generate_document(self, lease_data: Dict[str, Any], output_path: str) -> str:
        """
        Génère un document de bail
        
        Args:
            lease_data: Données du bail
            output_path: Chemin de sortie du document
            
        Returns:
            Chemin du fichier généré
        """
        pass
    
    @abstractmethod
    def get_supported_format(self) -> str:
        """Retourne le format supporté par ce générateur"""
        pass
    
    @abstractmethod
    def validate_data(self, lease_data: Dict[str, Any]) -> bool:
        """Valide que les données sont complètes pour la génération"""
        pass


class PDFLeaseGenerator(ILeaseDocumentGenerator):
    """Générateur de baux au format PDF avec ReportLab"""
    
    def __init__(self):
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab n'est pas installé. Exécutez: pip install reportlab")
        
        self.styles = getSampleStyleSheet()
        self._setup_styles()
    
    def _setup_styles(self):
        """Configure les styles personnalisés"""
        
        # Style pour le titre principal
        self.styles.add(ParagraphStyle(
            name='LeaseTitle',
            parent=self.styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=black
        ))
        
        # Style pour les sections
        self.styles.add(ParagraphStyle(
            name='SectionTitle',
            parent=self.styles['Heading2'],
            fontSize=14,
            spaceBefore=20,
            spaceAfter=12,
            textColor=blue
        ))
        
        # Style pour le contenu
        self.styles.add(ParagraphStyle(
            name='LeaseContent',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            alignment=TA_JUSTIFY,
            firstLineIndent=0.5*cm
        ))
        
        # Style pour les signatures
        self.styles.add(ParagraphStyle(
            name='Signature',
            parent=self.styles['Normal'],
            fontSize=10,
            spaceBefore=30,
            alignment=TA_LEFT
        ))
    
    def get_supported_format(self) -> str:
        return "pdf"
    
    def validate_data(self, lease_data: Dict[str, Any]) -> bool:
        """Valide les données requises pour la génération"""
        required_fields = [
            'title', 'lease_type', 'landlord_info', 
            'financial_terms', 'start_date'
        ]
        
        for field in required_fields:
            if field not in lease_data:
                return False
        
        # Vérifier les informations du propriétaire
        landlord_required = ['first_name', 'last_name', 'email']
        if not all(field in lease_data['landlord_info'] for field in landlord_required):
            return False
        
        # Vérifier les conditions financières
        financial_required = ['monthly_rent', 'charges', 'deposit']
        if not all(field in lease_data['financial_terms'] for field in financial_required):
            return False
        
        return True
    
    def generate_document(self, lease_data: Dict[str, Any], output_path: str) -> str:
        """Génère un PDF de bail"""
        
        if not self.validate_data(lease_data):
            raise ValueError("Données de bail incomplètes pour la génération")
        
        # Créer le répertoire de sortie si nécessaire
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Créer le document PDF
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        # Construire le contenu
        story = []
        
        # Titre principal
        title = lease_data.get('title', 'CONTRAT DE BAIL')
        story.append(Paragraph(title, self.styles['LeaseTitle']))
        story.append(Spacer(1, 20))
        
        # Informations générales
        self._add_general_info(story, lease_data)
        
        # Parties contractantes
        self._add_parties_section(story, lease_data)
        
        # Désignation du logement
        self._add_property_section(story, lease_data)
        
        # Conditions financières
        self._add_financial_section(story, lease_data)
        
        # Durée du bail
        self._add_duration_section(story, lease_data)
        
        # Clauses spécifiques
        self._add_clauses_section(story, lease_data)
        
        # Mobilier (si meublé)
        if lease_data.get('lease_type') in ['FURNISHED', 'STUDENT']:
            self._add_furniture_section(story, lease_data)
        
        # Signatures
        self._add_signatures_section(story, lease_data)
        
        # Générer le PDF
        doc.build(story)
        
        return output_path
    
    def _add_general_info(self, story: List, lease_data: Dict[str, Any]):
        """Ajoute les informations générales"""
        story.append(Paragraph("INFORMATIONS GÉNÉRALES", self.styles['SectionTitle']))
        
        lease_type_labels = {
            'UNFURNISHED': 'Location vide (loi du 6 juillet 1989)',
            'FURNISHED': 'Location meublée',
            'STUDENT': 'Location étudiante',
            'SEASONAL': 'Location saisonnière',
            'COMMERCIAL': 'Location commerciale'
        }
        
        lease_type = lease_data.get('lease_type', 'UNFURNISHED')
        type_label = lease_type_labels.get(lease_type, lease_type)
        
        content = f"""
        <b>Type de bail :</b> {type_label}<br/>
        <b>Date d'établissement :</b> {datetime.now().strftime('%d/%m/%Y')}<br/>
        <b>Référence :</b> BAIL-{lease_data.get('id', 'XXX')}-{datetime.now().year}
        """
        
        story.append(Paragraph(content, self.styles['LeaseContent']))
        story.append(Spacer(1, 15))
    
    def _add_parties_section(self, story: List, lease_data: Dict[str, Any]):
        """Ajoute la section des parties contractantes"""
        story.append(Paragraph("DÉSIGNATION DES PARTIES", self.styles['SectionTitle']))
        
        # Propriétaire/Bailleur
        landlord = lease_data['landlord_info']
        landlord_content = f"""
        <b>LE BAILLEUR :</b><br/>
        {landlord.get('first_name', '')} {landlord.get('last_name', '')}<br/>
        """
        
        if landlord.get('address'):
            landlord_content += f"Adresse : {landlord['address']}<br/>"
        if landlord.get('postal_code') and landlord.get('city'):
            landlord_content += f"{landlord['postal_code']} {landlord['city']}<br/>"
        if landlord.get('email'):
            landlord_content += f"Email : {landlord['email']}<br/>"
        if landlord.get('phone'):
            landlord_content += f"Téléphone : {landlord['phone']}<br/>"
        
        story.append(Paragraph(landlord_content, self.styles['LeaseContent']))
        
        # Locataire (si défini)
        if lease_data.get('tenant_info'):
            tenant = lease_data['tenant_info']
            tenant_content = f"""
            <b>LE LOCATAIRE :</b><br/>
            {tenant.get('first_name', '')} {tenant.get('last_name', '')}<br/>
            """
            
            if tenant.get('address'):
                tenant_content += f"Adresse : {tenant['address']}<br/>"
            if tenant.get('postal_code') and tenant.get('city'):
                tenant_content += f"{tenant['postal_code']} {tenant['city']}<br/>"
            if tenant.get('email'):
                tenant_content += f"Email : {tenant['email']}<br/>"
            if tenant.get('phone'):
                tenant_content += f"Téléphone : {tenant['phone']}<br/>"
            
            story.append(Paragraph(tenant_content, self.styles['LeaseContent']))
        else:
            story.append(Paragraph("<b>LE LOCATAIRE :</b> À compléter", self.styles['LeaseContent']))
        
        # Garant (si défini)
        if lease_data.get('guarantor_info'):
            guarantor = lease_data['guarantor_info']
            guarantor_content = f"""
            <b>LA CAUTION SOLIDAIRE :</b><br/>
            {guarantor.get('first_name', '')} {guarantor.get('last_name', '')}<br/>
            """
            
            if guarantor.get('email'):
                guarantor_content += f"Email : {guarantor['email']}<br/>"
            if guarantor.get('phone'):
                guarantor_content += f"Téléphone : {guarantor['phone']}<br/>"
            
            story.append(Paragraph(guarantor_content, self.styles['LeaseContent']))
        
        story.append(Spacer(1, 15))
    
    def _add_property_section(self, story: List, lease_data: Dict[str, Any]):
        """Ajoute la section de désignation du logement"""
        story.append(Paragraph("DÉSIGNATION DU LOGEMENT", self.styles['SectionTitle']))
        
        apartment_data = lease_data.get('apartment_data', {})
        
        property_content = f"""
        <b>Adresse du logement :</b><br/>
        {apartment_data.get('address', 'À compléter')}<br/>
        {apartment_data.get('postal_code', '')} {apartment_data.get('city', '')}<br/><br/>
        
        <b>Description :</b><br/>
        Type : {apartment_data.get('apartment_type', 'Appartement')}<br/>
        Surface : {apartment_data.get('surface_area', 'À compléter')} m²<br/>
        Nombre de pièces : {apartment_data.get('room_count', 'À compléter')}<br/>
        Étage : {apartment_data.get('floor', 'À compléter')}<br/>
        """
        
        if apartment_data.get('building_year'):
            property_content += f"Année de construction : {apartment_data['building_year']}<br/>"
        
        story.append(Paragraph(property_content, self.styles['LeaseContent']))
        story.append(Spacer(1, 15))
    
    def _add_financial_section(self, story: List, lease_data: Dict[str, Any]):
        """Ajoute la section des conditions financières"""
        story.append(Paragraph("CONDITIONS FINANCIÈRES", self.styles['SectionTitle']))
        
        financial = lease_data['financial_terms']
        
        # Création d'un tableau pour les montants
        financial_data = [
            ['Description', 'Montant'],
            ['Loyer mensuel hors charges', f"{financial['monthly_rent']:.2f} €"],
            ['Charges mensuelles', f"{financial['charges']:.2f} €"],
            ['Total mensuel TTC', f"{financial['monthly_rent'] + financial['charges']:.2f} €"],
            ['Dépôt de garantie', f"{financial['deposit']:.2f} €"]
        ]
        
        if financial.get('agency_fees'):
            financial_data.append(['Honoraires d\'agence', f"{financial['agency_fees']:.2f} €"])
        
        table = Table(financial_data, colWidths=[8*cm, 4*cm])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), blue),
            ('TEXTCOLOR', (0, 0), (-1, 0), black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), '#f0f0f0'),
            ('GRID', (0, 0), (-1, -1), 1, black)
        ]))
        
        story.append(table)
        
        # Informations sur le paiement
        payment_info = f"""
        <b>Modalités de paiement :</b><br/>
        Le loyer est payable d'avance le {financial.get('rent_payment_day', 1)} de chaque mois.<br/>
        Indexation sur l'indice : {financial.get('indexation_reference', 'IRL')}
        """
        
        story.append(Spacer(1, 10))
        story.append(Paragraph(payment_info, self.styles['LeaseContent']))
        story.append(Spacer(1, 15))
    
    def _add_duration_section(self, story: List, lease_data: Dict[str, Any]):
        """Ajoute la section durée du bail"""
        story.append(Paragraph("DURÉE DU BAIL", self.styles['SectionTitle']))
        
        start_date = lease_data.get('start_date')
        duration_months = lease_data.get('duration_months', 36)
        
        if isinstance(start_date, str):
            start_date = datetime.fromisoformat(start_date).date()
        elif isinstance(start_date, datetime):
            start_date = start_date.date()
        
        # Calculer la date de fin
        end_year = start_date.year + (duration_months // 12)
        end_month = start_date.month + (duration_months % 12)
        if end_month > 12:
            end_year += 1
            end_month -= 12
        
        try:
            end_date = start_date.replace(year=end_year, month=end_month)
        except ValueError:
            # Gérer le cas du 29 février dans une année non bissextile
            end_date = start_date.replace(year=end_year, month=end_month, day=28)
        
        duration_content = f"""
        <b>Date d'entrée en vigueur :</b> {start_date.strftime('%d/%m/%Y')}<br/>
        <b>Durée :</b> {duration_months} mois<br/>
        <b>Date d'échéance :</b> {end_date.strftime('%d/%m/%Y')}<br/><br/>
        
        Le présent bail est conclu pour une durée de {duration_months} mois.
        """
        
        # Ajouter les conditions de renouvellement selon le type de bail
        lease_type = lease_data.get('lease_type')
        if lease_type == 'UNFURNISHED':
            duration_content += """<br/><br/>
            Conformément à la loi du 6 juillet 1989, ce bail se renouvellera automatiquement 
            par tacite reconduction pour une durée de 3 ans, sauf congé donné par l'une des parties 
            dans les conditions prévues par la loi.
            """
        elif lease_type == 'FURNISHED':
            duration_content += """<br/><br/>
            Ce bail de location meublée se renouvellera automatiquement par tacite reconduction 
            pour une durée de 1 an, sauf congé donné par l'une des parties.
            """
        
        story.append(Paragraph(duration_content, self.styles['LeaseContent']))
        story.append(Spacer(1, 15))
    
    def _add_clauses_section(self, story: List, lease_data: Dict[str, Any]):
        """Ajoute les clauses du bail"""
        story.append(Paragraph("CLAUSES PARTICULIÈRES", self.styles['SectionTitle']))
        
        # Clauses standard selon le type
        standard_clauses = self._get_standard_clauses(lease_data.get('lease_type'))
        
        for i, clause in enumerate(standard_clauses, 1):
            clause_content = f"<b>Article {i} - {clause['title']}</b><br/>{clause['content']}"
            story.append(Paragraph(clause_content, self.styles['LeaseContent']))
            story.append(Spacer(1, 10))
        
        # Clauses personnalisées
        custom_clauses = lease_data.get('custom_clauses', [])
        for i, clause in enumerate(custom_clauses, len(standard_clauses) + 1):
            clause_content = f"<b>Article {i} - Clause personnalisée</b><br/>{clause}"
            story.append(Paragraph(clause_content, self.styles['LeaseContent']))
            story.append(Spacer(1, 10))
        
        # Notes additionnelles
        if lease_data.get('additional_notes'):
            story.append(Paragraph("<b>Notes additionnelles :</b>", self.styles['SectionTitle']))
            story.append(Paragraph(lease_data['additional_notes'], self.styles['LeaseContent']))
        
        story.append(Spacer(1, 15))
    
    def _add_furniture_section(self, story: List, lease_data: Dict[str, Any]):
        """Ajoute l'inventaire du mobilier pour les baux meublés"""
        story.append(Paragraph("INVENTAIRE DU MOBILIER", self.styles['SectionTitle']))
        
        furniture_inventory = lease_data.get('furniture_inventory', [])
        
        if furniture_inventory:
            # Créer un tableau pour l'inventaire
            furniture_data = [['Élément', 'Catégorie', 'Quantité', 'État', 'Valeur estimée']]
            
            total_value = 0
            for item in furniture_inventory:
                value = item.get('estimated_value', 0) or 0
                total_value += value * item.get('quantity', 1)
                
                furniture_data.append([
                    item.get('name', ''),
                    item.get('category', ''),
                    str(item.get('quantity', 1)),
                    item.get('condition', 'Bon'),
                    f"{value:.2f} €" if value else "N/A"
                ])
            
            # Ligne de total
            furniture_data.append(['', '', '', 'TOTAL', f"{total_value:.2f} €"])
            
            table = Table(furniture_data, colWidths=[4*cm, 3*cm, 2*cm, 2*cm, 3*cm])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), blue),
                ('TEXTCOLOR', (0, 0), (-1, 0), black),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BACKGROUND', (0, 1), (-1, -2), '#f9f9f9'),
                ('BACKGROUND', (0, -1), (-1, -1), '#e0e0e0'),
                ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 1, black)
            ]))
            
            story.append(table)
        else:
            story.append(Paragraph("Inventaire du mobilier à compléter lors de l'état des lieux d'entrée.", 
                                 self.styles['LeaseContent']))
        
        story.append(Spacer(1, 15))
    
    def _add_signatures_section(self, story: List, lease_data: Dict[str, Any]):
        """Ajoute la section des signatures"""
        story.append(Paragraph("SIGNATURES", self.styles['SectionTitle']))
        
        signature_content = f"""
        Fait à ________________________, le _______________<br/><br/>
        
        <b>Le Bailleur :</b><br/>
        {lease_data['landlord_info'].get('first_name', '')} {lease_data['landlord_info'].get('last_name', '')}<br/>
        Lu et approuvé,<br/><br/>
        Signature :<br/><br/><br/><br/>
        
        <b>Le Locataire :</b><br/>
        """
        
        if lease_data.get('tenant_info'):
            tenant = lease_data['tenant_info']
            signature_content += f"{tenant.get('first_name', '')} {tenant.get('last_name', '')}<br/>"
        else:
            signature_content += "Nom et prénom : ________________________<br/>"
        
        signature_content += """
        Lu et approuvé,<br/><br/>
        Signature :<br/><br/><br/><br/>
        """
        
        # Garant si nécessaire
        if lease_data.get('guarantor_info'):
            guarantor = lease_data['guarantor_info']
            signature_content += f"""
            <b>La Caution Solidaire :</b><br/>
            {guarantor.get('first_name', '')} {guarantor.get('last_name', '')}<br/>
            Lu et approuvé,<br/><br/>
            Signature :<br/><br/><br/><br/>
            """
        
        story.append(Paragraph(signature_content, self.styles['Signature']))
    
    def _get_standard_clauses(self, lease_type: str) -> List[Dict[str, str]]:
        """Retourne les clauses standard selon le type de bail"""
        
        common_clauses = [
            {
                "title": "Destination des lieux",
                "content": "Le logement est loué à usage d'habitation principale exclusivement."
            },
            {
                "title": "État des lieux",
                "content": "Un état des lieux contradictoire sera établi lors de la remise des clés et devra être joint au présent contrat."
            },
            {
                "title": "Obligations du locataire",
                "content": "Le locataire s'engage à occuper paisiblement les lieux, à les entretenir en bon père de famille, à ne pas transformer les lieux sans accord écrit du bailleur, et à s'acquitter du loyer et des charges aux échéances convenues."
            },
            {
                "title": "Obligations du bailleur",
                "content": "Le bailleur s'engage à délivrer un logement décent et en bon état d'usage et de réparation, à assurer la jouissance paisible du logement et à effectuer les réparations autres que locatives."
            },
            {
                "title": "Résiliation",
                "content": "Le locataire peut résilier le bail à tout moment en respectant un préavis de 3 mois (1 mois pour les logements situés en zone tendue, les logements meublés, ou en cas de mutation professionnelle)."
            }
        ]
        
        if lease_type == 'FURNISHED':
            common_clauses.append({
                "title": "Mobilier et équipements",
                "content": "Le logement est loué avec le mobilier et les équipements décrits dans l'inventaire annexé au présent contrat. Le locataire s'engage à en prendre soin et à les restituer en bon état."
            })
        
        return common_clauses


class WordLeaseGenerator(ILeaseDocumentGenerator):
    """Générateur de baux au format Word (optionnel)"""
    
    def __init__(self):
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx n'est pas installé. Exécutez: pip install python-docx")
    
    def get_supported_format(self) -> str:
        return "docx"
    
    def validate_data(self, lease_data: Dict[str, Any]) -> bool:
        # Même validation que pour PDF
        return PDFLeaseGenerator().validate_data(lease_data)
    
    def generate_document(self, lease_data: Dict[str, Any], output_path: str) -> str:
        """Génère un document Word de bail"""
        
        if not self.validate_data(lease_data):
            raise ValueError("Données de bail incomplètes pour la génération")
        
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        doc = Document()
        
        # Titre
        title = doc.add_heading(lease_data.get('title', 'CONTRAT DE BAIL'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ajouter les sections (implémentation simplifiée)
        doc.add_heading('DÉSIGNATION DES PARTIES', level=1)
        
        landlord = lease_data['landlord_info']
        doc.add_paragraph(f"LE BAILLEUR : {landlord.get('first_name', '')} {landlord.get('last_name', '')}")
        
        if lease_data.get('tenant_info'):
            tenant = lease_data['tenant_info']
            doc.add_paragraph(f"LE LOCATAIRE : {tenant.get('first_name', '')} {tenant.get('last_name', '')}")
        
        # Conditions financières
        doc.add_heading('CONDITIONS FINANCIÈRES', level=1)
        financial = lease_data['financial_terms']
        doc.add_paragraph(f"Loyer mensuel : {financial['monthly_rent']:.2f} €")
        doc.add_paragraph(f"Charges : {financial['charges']:.2f} €")
        doc.add_paragraph(f"Dépôt de garantie : {financial['deposit']:.2f} €")
        
        doc.save(output_path)
        return output_path


class LeaseDocumentService:
    """Service principal pour la génération de documents de bail"""
    
    def __init__(self):
        self.generators = {}
        self._register_generators()
    
    def _register_generators(self):
        """Enregistre les générateurs disponibles"""
        
        if REPORTLAB_AVAILABLE:
            self.generators['pdf'] = PDFLeaseGenerator()
        
        if PYTHON_DOCX_AVAILABLE:
            self.generators['docx'] = WordLeaseGenerator()
    
    def get_available_formats(self) -> List[str]:
        """Retourne les formats disponibles"""
        return list(self.generators.keys())
    
    def generate_lease_document(
        self, 
        lease_data: Dict[str, Any], 
        format: str = 'pdf',
        output_dir: str = './generated_leases'
    ) -> Dict[str, Any]:
        """
        Génère un document de bail
        
        Args:
            lease_data: Données du bail
            format: Format souhaité (pdf, docx)
            output_dir: Répertoire de sortie
            
        Returns:
            Informations sur le document généré
        """
        
        if format not in self.generators:
            available = ', '.join(self.get_available_formats())
            raise ValueError(f"Format '{format}' non supporté. Formats disponibles: {available}")
        
        generator = self.generators[format]
        
        # Générer un nom de fichier unique
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        lease_id = lease_data.get('id', 'new')
        filename = f"bail_{lease_id}_{timestamp}.{format}"
        output_path = os.path.join(output_dir, filename)
        
        # Générer le document
        generated_path = generator.generate_document(lease_data, output_path)
        
        # Calculer le hash et la taille
        with open(generated_path, 'rb') as f:
            content = f.read()
            file_hash = hashlib.sha256(content).hexdigest()
            file_size = len(content)
        
        return {
            'file_path': generated_path,
            'filename': filename,
            'format': format,
            'file_size': file_size,
            'file_hash': file_hash,
            'generated_at': datetime.now(),
            'generator': generator.__class__.__name__
        }
    
    def validate_lease_data(self, lease_data: Dict[str, Any], format: str = 'pdf') -> bool:
        """Valide les données pour un format donné"""
        
        if format not in self.generators:
            return False
        
        return self.generators[format].validate_data(lease_data)


# Factory pour créer le service
def create_lease_document_service() -> LeaseDocumentService:
    """Factory pour créer le service de génération de documents"""
    return LeaseDocumentService()